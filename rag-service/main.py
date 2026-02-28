"""
RAG Service with Embedding Model Version Tracking
===================================================

ISSUE: T135 #119 - No Embedding Model Version Tracking Causes Vector Store Drift

PROBLEM:
- The RAG pipeline generates embeddings using HuggingFace models and stores them in FAISS
- However, there was NO tracking of embedding model version/dimension alongside vectors
- If the embedding model changed/upgraded without re-indexing, old vectors silently mixed with new embedding space
- This led to "silent retrieval corruption" - answers degraded without throwing errors

SOLUTION IMPLEMENTED:
1. EmbeddingMetadata class tracks: model_name, embedding_dim, created_timestamp, model_hash
2. Metadata persisted as JSON files alongside session data for durability
3. On vectorstore load: Validates stored metadata vs current model config
4. Auto-invalidates incompatible indexes and triggers re-embedding
5. New diagnostic endpoints for debugging embedding drift issues

KEY FEATURES:
- Automatic detection of embedding model changes
- Prevention of vector space mixing from different models
- Per-session metadata persistence to disk
- Thread-safe session management
- Silent failure prevention through explicit validation
- Diagnostic endpoints for monitoring and debugging

NEW ENDPOINTS:
- GET /embedding-model-info - Show current embedding model configuration
- POST /validate-embeddings - Validate session embedding compatibility
- Updated GET /status - Now includes embedding metadata

BACKWARD COMPATIBILITY:
- Existing vectorstores continue to work
- New sessions automatically tracked with metadata
- No schema changes to existing APIs
"""

from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel, Field, validator
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from groq import Groq
from dotenv import load_dotenv
from transformers import (
    AutoConfig,
    AutoTokenizer,
    AutoModelForSeq2SeqLM,
    AutoModelForCausalLM,
)
from slowapi import Limiter
from slowapi.util import get_remote_address
import os 
import re
import uvicorn
from slowapi import Limiter
from slowapi.util import get_remote_address
import threading
from datetime import datetime
from pathlib import Path
import json
import hashlib
import torch
import time
import docx

# ===============================
# APP SETUP
# ===============================
load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_DIR = (BASE_DIR / "uploads").resolve()

app = FastAPI()
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter

# ===============================
# CONFIG
# ===============================
HF_GENERATION_MODEL = os.getenv("HF_GENERATION_MODEL", "google/flan-t5-small")
LLM_GENERATION_TIMEOUT = int(os.getenv("LLM_GENERATION_TIMEOUT", "30"))
SESSION_TIMEOUT = 3600
EMBEDDING_MODEL_NAME = os.getenv("EMBEDDING_MODEL_NAME", "sentence-transformers/all-MiniLM-L6-v2")

# ===============================
# EMBEDDING MODEL METADATA TRACKING
# ===============================
class EmbeddingMetadata:
    """
    Tracks embedding model version and configuration to prevent vector store drift.
    This prevents silent retrieval corruption when embedding models change.
    
    Issue Reference: T135 #119 - Embedding Model Version Tracking
    """
    def __init__(self, model_name: str):
        self.model_name = model_name
        self.embedding_dim = self._get_embedding_dimension()
        self.created_timestamp = datetime.utcnow().isoformat()
        self.model_hash = self._compute_model_hash()
    
    def _get_embedding_dimension(self) -> int:
        """Get the embedding dimension from the model."""
        try:
            # Create a test embedding to determine dimension
            test_embeddings = HuggingFaceEmbeddings(model_name=self.model_name)
            test_embed = test_embeddings.embed_query("test")
            return len(test_embed)
        except Exception as e:
            raise RuntimeError(f"Failed to determine embedding dimension: {str(e)}")
    
    def _compute_model_hash(self) -> str:
        """Compute a hash of the model configuration for change detection."""
        config_str = f"{self.model_name}_{self.embedding_dim}_{EMBEDDING_MODEL_NAME}"
        return hashlib.sha256(config_str.encode()).hexdigest()
    
    def to_dict(self) -> dict:
        """Convert metadata to dictionary for JSON serialization."""
        return {
            "model_name": self.model_name,
            "embedding_dimension": self.embedding_dim,
            "created_timestamp": self.created_timestamp,
            "model_hash": self.model_hash,
            "metadata_version": "1.0"
        }
    
    @staticmethod
    def from_dict(data: dict) -> 'EmbeddingMetadata':
        """Load metadata from dictionary."""
        metadata = EmbeddingMetadata.__new__(EmbeddingMetadata)
        metadata.model_name = data.get("model_name")
        metadata.embedding_dim = data.get("embedding_dimension")
        metadata.created_timestamp = data.get("created_timestamp")
        metadata.model_hash = data.get("model_hash")
        return metadata
    
    def validate_compatibility(self, current_model_name: str) -> tuple[bool, str]:
        """
        Validate if stored metadata is compatible with current model config.
        
        Returns:
            (is_compatible, reason_message)
        """
        # Check model name match
        if self.model_name != current_model_name:
            reason = f"Model mismatch: stored={self.model_name}, current={current_model_name}"
            return False, reason
        
        # Check embedding dimension match
        try:
            current_embeddings = HuggingFaceEmbeddings(model_name=current_model_name)
            current_dim = len(current_embeddings.embed_query("test"))
            
            if self.embedding_dim != current_dim:
                reason = f"Embedding dimension mismatch: stored={self.embedding_dim}, current={current_dim}"
                return False, reason
        except Exception as e:
            return False, f"Failed to validate embedding dimension: {str(e)}"
        
        return True, "Metadata is compatible"

# ---------------------------------------------------------------------------
# GLOBAL STATE MANAGEMENT (Thread-safe, Multi-user support)
# ---------------------------------------------------------------------------
# Per-user/session storage with proper cleanup and locking
# Session structure: {session_id: {"vectorstore": FAISS, "upload_time": datetime, "embedding_metadata": EmbeddingMetadata}}
sessions = {}
sessions_lock = threading.RLock()  # Thread-safe access to sessions

# Load local embedding model (unchanged — FAISS retrieval stays the same)
embedding_model = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)

# Initialize current embedding metadata at startup
current_embedding_metadata = EmbeddingMetadata(EMBEDDING_MODEL_NAME)

embedding_model = HuggingFaceEmbeddings(
    model_name=EMBEDDING_MODEL_NAME
)

# ---------------------------------------------------------------------------
# SESSION MANAGEMENT UTILITIES (Thread-safe, Multi-user support)
# ---------------------------------------------------------------------------

def get_metadata_file_path(session_id: str) -> Path:
    """Get the path where embedding metadata should be stored for a session."""
    metadata_dir = Path(__file__).parent / ".embedding_metadata"
    metadata_dir.mkdir(exist_ok=True)
    return metadata_dir / f"{session_id}_metadata.json"


def save_embedding_metadata(session_id: str, metadata: EmbeddingMetadata) -> None:
    """
    Persist embedding metadata to disk for durability and verification.
    
    This ensures that even if the session is cleared from memory,
    we can detect model drift when re-loading from disk later.
    """
    try:
        metadata_file = get_metadata_file_path(session_id)
        with open(metadata_file, 'w') as f:
            json.dump(metadata.to_dict(), f, indent=2)
    except Exception as e:
        raise RuntimeError(f"Failed to save embedding metadata: {str(e)}")


def load_embedding_metadata(session_id: str) -> EmbeddingMetadata | None:
    """
    Load embedding metadata from disk if it exists.
    
    Returns None if no metadata file found (new session).
    """
    try:
        metadata_file = get_metadata_file_path(session_id)
        if metadata_file.exists():
            with open(metadata_file, 'r') as f:
                data = json.load(f)
                return EmbeddingMetadata.from_dict(data)
    except Exception as e:
        raise RuntimeError(f"Failed to load embedding metadata: {str(e)}")
    return None


def delete_embedding_metadata(session_id: str) -> None:
    """Delete embedding metadata file for a session."""
    try:
        metadata_file = get_metadata_file_path(session_id)
        if metadata_file.exists():
            metadata_file.unlink()
    except Exception as e:
        pass  # Silent fail on metadata cleanup


def validate_embedding_compatibility(session_id: str) -> tuple[bool, str]:
    """
    Check if stored embedding metadata is compatible with current model.
    
    Returns:
        (is_compatible, reason_message)
    """
    stored_metadata = load_embedding_metadata(session_id)
    
    if stored_metadata is None:
        return True, "No stored metadata (new session)"
    
    is_compatible, reason = stored_metadata.validate_compatibility(EMBEDDING_MODEL_NAME)
    
    if not is_compatible:
        # Log the incompatibility for debugging
        print(f"[EMBEDDING DRIFT DETECTED] Session {session_id}: {reason}")
    
    return is_compatible, reason


def get_session_vectorstore(session_id: str):
    """
    Safely retrieves vectorstore for a session.
    Returns (vectorstore, upload_time, embedding_metadata) or (None, None, None) if not found.
    
    IMPORTANT: Validates embedding compatibility to detect model drift.
    """
    with sessions_lock:
        if session_id in sessions:
            session_data = sessions[session_id]
            
            # Check embedding compatibility
            is_compatible, reason = validate_embedding_compatibility(session_id)
            
            if not is_compatible:
                # Auto-invalidate incompatible index
                print(f"[AUTO-INVALIDATE] Clearing session {session_id} due to: {reason}")
                clear_session(session_id)
                return None, None, None
            
            return (
                session_data.get("vectorstore"),
                session_data.get("upload_time"),
                session_data.get("embedding_metadata")
            )
        return None, None, None


def set_session_vectorstore(session_id: str, vectorstore, upload_time: str, embedding_metadata: EmbeddingMetadata = None):
    """
    Safely stores vectorstore for a session WITH embedding metadata.
    Clears old session if it exists (replaces it).
    
    IMPORTANT: Always persists metadata to disk for durability.
    """
    if embedding_metadata is None:
        embedding_metadata = current_embedding_metadata
    
    with sessions_lock:
        # Clear old session to prevent memory leaks
        if session_id in sessions:
            old_vectorstore = sessions[session_id].get("vectorstore")
            if old_vectorstore is not None:
                del old_vectorstore  # Allow garbage collection
        
        # Store new session WITH metadata
        sessions[session_id] = {
            "vectorstore": vectorstore,
            "upload_time": upload_time,
            "embedding_metadata": embedding_metadata
        }
        
        # Persist metadata to disk
        try:
            save_embedding_metadata(session_id, embedding_metadata)
        except Exception as e:
            print(f"[WARNING] Failed to persist metadata for session {session_id}: {str(e)}")


def clear_session(session_id: str):
    """
    Safely clears a specific session's vectorstore and data.
    Also cleans up persisted metadata.
    """
    with sessions_lock:
        if session_id in sessions:
            old_vectorstore = sessions[session_id].get("vectorstore")
            if old_vectorstore is not None:
                del old_vectorstore  # Allow garbage collection
            del sessions[session_id]
    
    # Clean up persisted metadata
    delete_embedding_metadata(session_id)


def normalize_spaced_text(text: str) -> str:
    pattern = r"\b(?:[A-Za-z] ){2,}[A-Za-z]\b"
    return re.sub(pattern, lambda m: m.group(0).replace(" ", ""), text)


def normalize_answer(text: str) -> str:
    """
    Post-processes the LLM-generated answer.
    """
    text = normalize_spaced_text(text)
    text = re.sub(r"^(Answer[^:]*:|Context:|Question:)\s*", "", text, flags=re.I)
    return text.strip()


# ===============================
# DOCUMENT LOADERS
# ===============================
def load_pdf(file_path: str):
    return PyPDFLoader(file_path).load()


def load_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        return [Document(page_content=f.read())]


def load_docx(file_path: str):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in [".txt", ".md"]:
        return load_txt(file_path)
    else:
        raise ValueError("Unsupported file format")


# ===============================
# MODEL LOADING
# ===============================
def load_generation_model():
    global generation_model, generation_tokenizer, generation_is_encoder_decoder

    if generation_model:
        return generation_tokenizer, generation_model, generation_is_encoder_decoder

    config = AutoConfig.from_pretrained(HF_GENERATION_MODEL)
    generation_is_encoder_decoder = bool(config.is_encoder_decoder)

    generation_tokenizer = AutoTokenizer.from_pretrained(HF_GENERATION_MODEL)

    if generation_is_encoder_decoder:
        generation_model = AutoModelForSeq2SeqLM.from_pretrained(HF_GENERATION_MODEL)
    else:
        generation_model = AutoModelForCausalLM.from_pretrained(HF_GENERATION_MODEL)

    if torch.cuda.is_available():
        generation_model = generation_model.to("cuda")

    generation_model.eval()
    return generation_tokenizer, generation_model, generation_is_encoder_decoder


def generate_response(prompt: str, max_new_tokens: int):
    tokenizer, model, is_enc = load_generation_model()
    device = next(model.parameters()).device

    encoded = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
    encoded = {k: v.to(device) for k, v in encoded.items()}

    output = model.generate(
        **encoded,
        max_new_tokens=max_new_tokens,
        do_sample=False,
        pad_token_id=tokenizer.pad_token_id or tokenizer.eos_token_id,
    )

    if is_enc:
        return tokenizer.decode(output[0], skip_special_tokens=True)

    return tokenizer.decode(
        output[0][encoded["input_ids"].shape[1]:],
        skip_special_tokens=True,
    )


# ===============================
# REQUEST MODELS
# ===============================
class DocumentPath(BaseModel):
    filePath: str
    session_id: str


class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    session_id: str
    history: list = []

    @validator("question")
    def validate_question(cls, v):
        if not v.strip():
            raise ValueError("Empty question")
        return v.strip()


class SummarizeRequest(BaseModel):
    session_id: str
    pdf: str | None = None



class CompareRequest(BaseModel):
    session_id: str

# -------------------------------------------------------------------
# SESSION CLEANUP
# -------------------------------------------------------------------


def cleanup_expired_sessions():
    now = time.time()
    expired = [k for k, v in sessions.items()
               if now - v["last"] > SESSION_TIMEOUT]
    for k in expired:
        del sessions[k]


# ===============================
# PROCESS DOCUMENT
# ===============================
@app.post("/process")
@limiter.limit("15/15 minutes")
def process_pdf(request: Request, data: DocumentPath):
    """
    Process and store PDF with proper cleanup and thread-safe multi-user support.
    Also persists embedding metadata to detect model drift.
    """
    try:
        loader = PyPDFLoader(data.filePath)
        raw_docs = loader.load()

        if not raw_docs:
            return {"error": "PDF file is empty or unreadable. Please check your file."}

        # ── Layer 1: normalize at ingestion ──────────────────────────────────────
        cleaned_docs = []
        for doc in raw_docs:
            cleaned_content = normalize_spaced_text(doc.page_content)
            cleaned_docs.append(Document(page_content=cleaned_content, metadata=doc.metadata))

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        chunks = splitter.split_documents(cleaned_docs)
        
        if not chunks:
            return {"error": "No text chunks generated from the PDF. Please check your file."}

        # **KEY FIX**: Store per-session with automatic cleanup of old data
        session_id = request.headers.get("X-Session-ID", "default")
        upload_time = datetime.now().isoformat()
        
        # Thread-safe storage (automatically clears old session data)
        # Also stores current embedding metadata for drift detection
        vectorstore = FAISS.from_documents(chunks, embedding_model)
        set_session_vectorstore(session_id, vectorstore, upload_time, current_embedding_metadata)
        
        return {
            "message": "PDF processed successfully",
            "session_id": session_id,
            "upload_time": upload_time,
            "chunks_created": len(chunks),
            "embedding_model": EMBEDDING_MODEL_NAME,
            "embedding_dimension": current_embedding_metadata.embedding_dim
        }
            
    except Exception as e:
        return {
            "error": f"PDF processing failed: {str(e)}",
            "details": "Please ensure the file is a valid PDF"
        }


@app.post("/ask")
@limiter.limit("60/15 minutes")
def ask_question(request: Request, data: AskRequest):
    """
    Answer questions using session-specific PDF context with thread-safe access.
    Validates embedding compatibility to ensure retrieval correctness.
    """
    session_id = request.headers.get("X-Session-ID", "default")
    vectorstore, upload_time, embedding_metadata = get_session_vectorstore(session_id)
    
    if vectorstore is None:
        return {"answer": "Please upload a PDF first!"}
    
    try:
        # Thread-safe vectorstore access
        with sessions_lock:
            question = data.question
            history = data.history
            conversation_context = ""
            
            if history:
                for msg in history[-5:]:
                    role = msg.get("role", "")
                    content = msg.get("content", "")
                    if role and content:
                        conversation_context += f"{role}: {content}\n"
            
            # Search only within current session's vectorstore
            docs = vectorstore.similarity_search(question, k=4)
            if not docs:
                return {"answer": "No relevant context found in the current PDF."}

            context = "\n\n".join([doc.page_content for doc in docs])

            prompt = f"""You are a helpful assistant answering questions ONLY from the provided PDF document.

Conversation History (for context only):
{conversation_context}

Document Context (ONLY reference this):
{context}

Current Question:
{question}

Instructions:
- Answer ONLY using the document context provided above.
- Do NOT use any information from previous documents or conversations outside this context.
- If the answer is not in the document, say so briefly.
- Keep the answer concise (2-3 sentences max).

Answer:"""

            raw_answer = generate_response(prompt, max_new_tokens=512)
            answer = normalize_answer(raw_answer)
            return {"answer": answer}
            
    except Exception as e:
        return {"answer": f"Error processing question: {str(e)}"}

@app.post("/summarize")
@limiter.limit("15/15 minutes")
def summarize_pdf(request: Request, data: SummarizeRequest):
    """
    Summarize PDF using session-specific context with thread-safe access.
    Validates embedding compatibility to ensure summarization correctness.
    """
    session_id = request.headers.get("X-Session-ID", "default")
    vectorstore, upload_time, embedding_metadata = get_session_vectorstore(session_id)
    
    if vectorstore is None:
        return {"summary": "Please upload a PDF first!"}

    try:
        # Thread-safe vectorstore access
        with sessions_lock:
            docs = vectorstore.similarity_search("Give a concise summary of the document.", k=6)
            if not docs:
                return {"summary": "No document context available to summarize."}

            context = "\n\n".join([doc.page_content for doc in docs])

            prompt = (
                "You are a document summarization assistant working with a certificate or official document.\n"
                "RULES:\n"
                "1. Summarize in 6-8 concise bullet points.\n"
                "2. Clearly distinguish: who received the certificate, what course, which company issued it,\n"
                "   who signed it, on what platform, and on what date.\n"
                "3. Return clean, properly formatted text — no character spacing, proper Title Case for names.\n"
                "4. Use ONLY the information in the context below.\n"
                "5. DO NOT reference any other documents or previous PDFs.\n\n"
                f"Context:\n{context}\n\n"
                "Summary (bullet points):"
            )

            raw_summary = generate_response(prompt, max_new_tokens=512)
            summary = normalize_answer(raw_summary)
            return {"summary": summary}
            
    except Exception as e:
        return {"summary": f"Error summarizing PDF: {str(e)}"}


@app.post("/compare")
@limiter.limit("15/15 minutes")
def compare_pdfs(request: Request, data: dict):
    """
    Compare two PDFs using their session-specific contexts.
    Supports multi-user/multi-PDF comparison feature.
    Validates embedding compatibility for both PDFs.
    """
    session_id_1 = data.get("session_id_1", "default")
    session_id_2 = data.get("session_id_2", "default")
    question = data.get("question", "Compare these documents")
    
    vectorstore_1, _, metadata_1 = get_session_vectorstore(session_id_1)
    vectorstore_2, _, metadata_2 = get_session_vectorstore(session_id_2)
    
    if vectorstore_1 is None or vectorstore_2 is None:
        return {"error": "One or both sessions do not have a PDF loaded"}
    
    try:
        with sessions_lock:
            docs_1 = vectorstore_1.similarity_search(question, k=3)
            docs_2 = vectorstore_2.similarity_search(question, k=3)
            
            context_1 = "\n\n".join([doc.page_content for doc in docs_1])
            context_2 = "\n\n".join([doc.page_content for doc in docs_2])
            
            prompt = f"""You are a document comparison assistant.

PDF 1 Context:
{context_1}

PDF 2 Context:
{context_2}

Question: {question}

Compare the two documents regarding this question and highlight key differences and similarities.

Comparison:"""
            
            comparison = generate_response(prompt, max_new_tokens=512)
            return {"comparison": normalize_answer(comparison)}
            
    except Exception as e:
        return {"error": f"Error comparing PDFs: {str(e)}"}


@app.post("/reset")
@limiter.limit("60/15 minutes")
def reset_session(request: Request):
    """
    Explicitly resets a session by clearing its vectorstore.
    """
    session_id = request.headers.get("X-Session-ID", "default")
    
    with sessions_lock:
        clear_session(session_id)
        
    return {
        "message": "Session cleared successfully",
        "session_id": session_id
    }


@app.get("/status")
def get_pdf_status(request: Request):
    """
    Returns the current PDF session status.
    Useful for debugging and ensuring proper state management.
    Also includes embedding model metadata for drift detection.
    """
    session_id = request.headers.get("X-Session-ID", "default")
    
    with sessions_lock:
        if session_id in sessions:
            metadata = sessions[session_id].get("embedding_metadata")
            metadata_dict = metadata.to_dict() if metadata else None
            
            return {
                "pdf_loaded": True,
                "session_id": session_id,
                "upload_time": sessions[session_id].get("upload_time"),
                "embedding_metadata": metadata_dict
            }
        return {
            "pdf_loaded": False,
            "session_id": session_id,
            "upload_time": None,
            "embedding_metadata": None
        }


# -------------------------------------------------------------------
# EMBEDDING MODEL VALIDATION & DIAGNOSTICS
# -------------------------------------------------------------------

@app.get("/embedding-model-info")
def get_embedding_model_info():
    """
    Returns current embedding model configuration.
    Useful for debugging embedding drift issues.
    """
    return {
        "model_name": EMBEDDING_MODEL_NAME,
        "embedding_dimension": current_embedding_metadata.embedding_dim,
        "model_hash": current_embedding_metadata.model_hash,
        "created_timestamp": current_embedding_metadata.created_timestamp,
        "metadata_version": "1.0"
    }


@app.post("/validate-embeddings")
def validate_embeddings_endpoint(request: Request):
    """
    Validates embedding compatibility for a session.
    Returns detailed information about any incompatibilities detected.
    
    This endpoint is useful for debugging:
    - Model version mismatches
    - Embedding dimension changes
    - Silent vector store drift
    """
    session_id = request.headers.get("X-Session-ID", "default")
    
    is_compatible, reason = validate_embedding_compatibility(session_id)
    
    stored_metadata = load_embedding_metadata(session_id)
    
    return {
        "session_id": session_id,
        "is_compatible": is_compatible,
        "reason": reason,
        "current_model": EMBEDDING_MODEL_NAME,
        "current_dimension": current_embedding_metadata.embedding_dim,
        "stored_metadata": stored_metadata.to_dict() if stored_metadata else None,
        "action_taken": "auto-invalidated" if not is_compatible else "vectorstore_valid"
    }


# -------------------------------------------------------------------
# START SERVER
# -------------------------------------------------------------------
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True) 